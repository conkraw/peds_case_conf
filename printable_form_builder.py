"""Printable worksheet export for the Pediatric Case Conference Builder."""

from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from case_schema import CASE_SLIDES

BLUE = "1E5082"
LIGHT_BLUE = "E2ECF6"
GRAY = "F2F2F2"
DARK_GRAY = "555555"
WHITE = "FFFFFF"
DOC_FONT = "Calibri"


def _safe_text(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_borders(cell, color: str = "000000", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _format_run(run, size: float = 9, bold: bool = False, color: str | None = None, italic: bool = False):
    run.font.name = DOC_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _write_cell(cell, text: str, size: float = 9, bold: bool = False, color: str | None = None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _format_run(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _limits_text(field: Dict[str, Any]) -> str:
    parts: List[str] = ["Required" if field.get("required") else "Optional"]
    if "max_words" in field:
        parts.append(f"max {field['max_words']} words")
    if "max_lines" in field:
        parts.append(f"max {field['max_lines']} lines")
    if "max_words_per_line" in field:
        parts.append(f"max {field['max_words_per_line']} words/line")
    if "max_rows" in field:
        parts.append(f"max {field['max_rows']} rows")
    return " | ".join(parts)


def _writing_lines_for_field(field: Dict[str, Any]) -> int:
    if field.get("type") == "text" or field.get("type") == "select":
        return 1
    if "max_lines" in field:
        return min(max(int(field["max_lines"]), 1), 8)
    max_words = int(field.get("max_words", 35) or 35)
    if max_words <= 25:
        return 1
    if max_words <= 45:
        return 2
    if max_words <= 70:
        return 3
    return 4


def _add_section_heading(doc: Document, label: str, title: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_borders(cell)
    _set_cell_shading(cell, BLUE)
    _write_cell(cell, f"{label} | {title}" if title and title != label else label, size=10.5, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_field_box(doc: Document, field: Dict[str, Any]) -> None:
    if field.get("type") == "table":
        columns = field.get("columns", [])
        header = doc.add_table(rows=1, cols=1)
        header.style = "Table Grid"
        hcell = header.cell(0, 0)
        _set_cell_borders(hcell)
        _set_cell_shading(hcell, GRAY)
        _write_cell(hcell, f"{field['label']} ({_limits_text(field)})", size=8.8, bold=True, color=BLUE)
        if field.get("guide"):
            p = hcell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(field["guide"])
            _format_run(r, size=7.6, color=DARK_GRAY, italic=True)
        rows = int(field.get("max_rows", 5)) + 1
        table = doc.add_table(rows=rows, cols=len(columns))
        table.style = "Table Grid"
        for row in table.rows:
            for cell in row.cells:
                _set_cell_borders(cell)
        for c, col in enumerate(columns):
            cell = table.cell(0, c)
            _set_cell_shading(cell, BLUE)
            _write_cell(cell, col, size=7.5, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        for r in range(1, rows):
            for c in range(len(columns)):
                _write_cell(table.cell(r, c), "\n", size=8)
        doc.add_paragraph()
        return

    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)
    header = table.cell(0, 0)
    _set_cell_shading(header, GRAY)
    _write_cell(header, f"{field['label']} ({_limits_text(field)})", size=8.8, bold=True, color=BLUE)
    if field.get("guide"):
        p = header.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(field["guide"])
        _format_run(r, size=7.5, color=DARK_GRAY, italic=True)
    body = table.cell(1, 0)
    _write_cell(body, "\n" * _writing_lines_for_field(field), size=9)
    doc.add_paragraph()


def build_printable_planning_form() -> BytesIO:
    """Create a printable worksheet that mirrors the Streamlit fields."""
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    doc.styles["Normal"].font.name = DOC_FONT
    doc.styles["Normal"].font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PEDIATRIC CASE CONFERENCE BUILDER")
    _format_run(r, size=16, bold=True, color=BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Printable Planning Worksheet")
    _format_run(r2, size=11, bold=True, color=DARK_GRAY)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    run = intro.add_run(
        "Use this worksheet to build a progressive, interactive case conference. "
        "The fields and limits match the Streamlit app. Transfer final answers into the app for export."
    )
    _format_run(run, size=8.5)

    for idx, slide in enumerate(CASE_SLIDES):
        if idx > 0:
            doc.add_page_break()
        _add_section_heading(doc, slide["label"], slide.get("export_title", ""))
        for field in slide.get("fields", []):
            _add_field_box(doc, field)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

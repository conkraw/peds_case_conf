"""Word export logic for the Pediatric Case Conference Builder."""

from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, Iterable, List

import qrcode
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from case_schema import CASE_SLIDES, archive_summary
from feedback_config import FEEDBACK_DISPLAY_URL, FEEDBACK_QR_URL

BLUE = "1E5082"
LIGHT_BLUE = "E2ECF6"
GRAY = "F2F2F2"
DARK_GRAY = "555555"
WHITE = "FFFFFF"
BORDER = "000000"
DOC_FONT = "Calibri"


def _safe_text(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _lines(value: Any, limit: int | None = None) -> List[str]:
    lines = [line.strip() for line in _safe_text(value).splitlines() if line.strip()]
    return lines[:limit] if limit else lines


def _bullet_text(items: Iterable[str], limit: int | None = None) -> str:
    cleaned = [str(item).strip() for item in items if str(item).strip()]
    if limit is not None:
        cleaned = cleaned[:limit]
    return "\n".join(f"• {item}" for item in cleaned)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_borders(cell, color: str = BORDER, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _format_run(run, size: float = 9, bold: bool = False, color: str | None = None, italic: bool = False):
    run.font.name = DOC_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _clear_cell(cell) -> None:
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph()


def _write_cell(cell, text: Any, size: float = 9, bold: bool = False, color: str | None = None, align=WD_ALIGN_PARAGRAPH.LEFT):
    _clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    parts = _safe_text(text).splitlines() or [""]
    for idx, part in enumerate(parts):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(part)
        _format_run(run, size=size, bold=bold, color=color)


def _style_table(table):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _add_heading_band(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    _style_table(table)
    cell = table.cell(0, 0)
    _shade_cell(cell, BLUE)
    _write_cell(cell, text, size=10.5, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_label_value_table(doc: Document, rows: List[tuple[str, str]], label_width: float = 1.7):
    table = doc.add_table(rows=len(rows), cols=2)
    _style_table(table)
    for i, (label, value) in enumerate(rows):
        label_cell = table.cell(i, 0)
        value_cell = table.cell(i, 1)
        _shade_cell(label_cell, GRAY)
        _write_cell(label_cell, label, size=8.8, bold=True, color=BLUE)
        _write_cell(value_cell, value, size=8.8)
    return table


def _add_data_table(doc: Document, rows: List[Dict[str, Any]], columns: List[str]):
    clean_rows = []
    for row in rows or []:
        if any(_safe_text(row.get(col, "")) for col in columns):
            clean_rows.append(row)
    clean_rows = clean_rows[:6]
    table = doc.add_table(rows=len(clean_rows) + 1, cols=len(columns))
    _style_table(table)
    for c, col in enumerate(columns):
        cell = table.cell(0, c)
        _shade_cell(cell, BLUE)
        _write_cell(cell, col, size=8, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(clean_rows, start=1):
        for c, col in enumerate(columns):
            cell = table.cell(r, c)
            if r % 2 == 0:
                _shade_cell(cell, GRAY)
            _write_cell(cell, _safe_text(row.get(col, "")), size=7.8)
    return table


def _setup_doc(doc: Document, landscape: bool = False):
    section = doc.sections[-1]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    styles = doc.styles
    styles["Normal"].font.name = DOC_FONT
    styles["Normal"].font.size = Pt(9)


def _title_paragraph(doc: Document, title: str, subtitle: str = ""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    _format_run(r, size=16, bold=True, color=BLUE)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(subtitle)
        _format_run(r2, size=9.5, color=DARK_GRAY)


def build_word_summary(deck: Dict[str, Dict[str, Any]]) -> BytesIO:
    """Build a compact case conference summary DOCX."""
    doc = Document()
    _setup_doc(doc, landscape=False)
    archive = archive_summary(deck)
    title = archive.get("case") or "Case Conference"
    presenter = archive.get("presenter") or "Presenter"
    _title_paragraph(doc, "Pediatric Case Conference Summary", f"{title} | {presenter}")

    title_data = deck.get("title_goal", {})
    problem = deck.get("problem_representation", {})
    differential = deck.get("differential", {})
    diagnostics = deck.get("diagnostic_interpretation", {})
    management = deck.get("management_decision", {})
    teaching = deck.get("disease_teaching", {})
    points = deck.get("major_learning_points", {})
    patient = deck.get("return_to_patient", {})
    bottom = deck.get("final_bottom_line", {})

    _add_heading_band(doc, "Session Focus")
    _add_label_value_table(
        doc,
        [
            ("Learning point", archive.get("learning_point", "")),
            ("Session goal", _safe_text(title_data.get("session_goal"))),
            ("Reasoning skill", _safe_text(title_data.get("reasoning_skill"))),
            ("Privacy check", _safe_text(title_data.get("privacy_check"))),
        ],
    )

    _add_heading_band(doc, "Clinical Reasoning")
    _add_label_value_table(
        doc,
        [
            ("Problem representation", _safe_text(problem.get("one_liner"))),
            ("Clinical pivot", _safe_text(problem.get("pivot_point"))),
            ("Must-not-miss", _bullet_text(_lines(differential.get("must_not_miss"), 3))),
            ("Diagnostic bottom line", _safe_text(diagnostics.get("diagnostic_bottom_line"))),
            ("Management", _safe_text(management.get("initial_management"))),
        ],
    )

    _add_heading_band(doc, "Disease Teaching")
    _add_label_value_table(
        doc,
        [
            ("Disease / syndrome", _safe_text(teaching.get("disease_name"))),
            ("Overview", _safe_text(teaching.get("disease_overview"))),
            ("Major points", _bullet_text(_lines(points.get("major_points"), 3))),
            ("Return to patient", _safe_text(patient.get("clinical_course"))),
            ("Final bottom line", _safe_text(bottom.get("final_bottom_line"))),
        ],
    )

    _add_heading_band(doc, "Feedback")
    table = doc.add_table(rows=1, cols=2)
    _style_table(table)
    _write_cell(table.cell(0, 0), f"Feedback link:\n{FEEDBACK_DISPLAY_URL}", size=8.5, bold=False)
    qr_img = BytesIO()
    qrcode.make(FEEDBACK_QR_URL).save(qr_img, format="PNG")
    qr_img.seek(0)
    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(qr_img, width=Inches(1.0), height=Inches(1.0))

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def build_review_text_docx(deck: Dict[str, Dict[str, Any]]) -> BytesIO:
    """Build a full mentor review DOCX containing every field."""
    doc = Document()
    _setup_doc(doc, landscape=False)
    archive = archive_summary(deck)
    _title_paragraph(doc, "Case Conference Mentor Review", f"{archive.get('case', '')} | {archive.get('presenter', '')}")

    for slide_def in CASE_SLIDES:
        _add_heading_band(doc, _safe_text(slide_def.get("export_title") or slide_def.get("label")))
        slide_data = deck.get(slide_def["id"], {})
        for field in slide_def.get("fields", []):
            value = slide_data.get(field["key"], "")
            if field.get("type") == "table":
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(field["label"])
                _format_run(r, size=9, bold=True, color=BLUE)
                _add_data_table(doc, value, field.get("columns", []))
            else:
                label = field["label"]
                if field.get("private_note"):
                    label += " (facilitator note)"
                _add_label_value_table(doc, [(label, _safe_text(value))])
        doc.add_paragraph()

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
